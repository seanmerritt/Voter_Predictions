import sklearn as sk
from imblearn.over_sampling import SMOTENC
import pandas as pd
import numpy as np
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.ensemble import GradientBoostingClassifier
from sklearn.ensemble import VotingClassifier, AdaBoostClassifier
from sklearn.model_selection import GridSearchCV, cross_val_score
from sklearn.decomposition import PCA
from sklearn.metrics import  roc_curve, auc, confusion_matrix, accuracy_score, roc_auc_score
import matplotlib.pylab as plt
from matplotlib.pylab import rcParams
rcParams['figure.figsize'] = 12, 4
from xgboost.sklearn import XGBClassifier
from docx import Document
from docx.text.paragraph import Paragraph
import os
from io import BytesIO

def modelfit(X_train, y_train, X_test, y_test, doc, title, performCV=True, printFeatureImportance=True, cv_folds=5):
    
    xgb = XGBClassifier( learning_rate =0.01, n_estimators= 2000, max_depth=6,
 min_child_weight=4, gamma=0, subsample=0.9, colsample_bytree=0.8,
 objective= 'binary:logistic', nthread=4, scale_pos_weight=1, seed=27, reg_alpha = .1)
    xgb.fit(X_train,y_train)

    y_preds = xgb.predict(X_test)
    dtrain_predprob = xgb.predict_proba(X_test)[:,1]
    
    #Print model report:
    #Print model report:
    doc.add_heading(f"\nModel Report {title}",0)
    doc.add_paragraph("Accuracy : %.4g" % accuracy_score(y_test, y_preds))
    doc.add_paragraph("AUC Score (Test): %f" % roc_auc_score(y_test, dtrain_predprob))

    #Perform cross-validation:
    if performCV:
        cv_score = cross_val_score(xgb, X_train, y_train, cv=cv_folds, scoring='roc_auc')
        doc.add_paragraph("CV Score : Mean - %.7g | Std - %.7g | Min - %.7g | Max - %.7g" % (np.mean(cv_score),np.std(cv_score),np.min(cv_score),np.max(cv_score)))   
     #Print Feature Importance:
    if printFeatureImportance:
        feat_plot = BytesIO()
        feat_imp = pd.Series(xgb.feature_importances_, X_train.columns).sort_values(ascending=False)
        feats = feat_imp.to_dict()
        for key in feats:
            doc.add_paragraph(f"{key}: {feats[key]}")
        feat_imp.plot(kind='bar', title='Feature Importances')
        plt.ylabel(f'Feature Importance Score for {title}')
        plt.savefig(feat_plot)

        doc.add_picture(feat_plot)
        feat_plot.close()

def prep_data(data, cols, cats):
    dat = data.dropna()
    y = dat['voted']
    X = dat.drop('voted', axis = 1)
    X = X.iloc[:,cols:]
    X, y = SMOTENC(categorical_features= cats).fit_resample(X,y)
    X_train, X_test, y_train, y_test = sk.model_selection.train_test_split(X,y, test_size= .25, random_state= 123)
    return X_train, X_test, y_train, y_test

def main():
    document = Document()
    os.chdir(r'C:\Users\seanm\OneDrive\Documents\C - Research\NonVoters\Voter_Predictions\Paper Drafts')

    # Both years
    dat = pd.read_csv(r'C:\Users\seanm\OneDrive\Documents\C - Research\NonVoters\Voter_Predictions\ERS_dat.csv')
    ## Without Voted previous and year 
    X_train, X_test, y_train, y_test = prep_data(dat, cols = 4, cats=[0,1,3,5,6,13])
    modelfit(X_train, y_train, X_test, y_test, document, title = '2012 & 2016 W/out previous vote', performCV=False, printFeatureImportance=True, cv_folds=5)
    ## With Voted previous and year 
    X_train, X_test, y_train, y_test = prep_data(dat, cols = 2,cats=[2,1,3,5,7,8,15])
    modelfit(X_train, y_train, X_test, y_test, document,  title= '2012 & 2016 W/ previous vote', performCV=False, printFeatureImportance=True, cv_folds=5)

    #2016
    dat = pd.read_csv(r'C:\Users\seanm\OneDrive\Documents\C - Research\NonVoters\Voter_Predictions\voting_2016_data.csv')
    dat = dat.drop('year', axis = 1)
    ## With Voted previous 
    X_train, X_test, y_train, y_test = prep_data(dat, cols = 2,  cats=[52,53,54,55,56,57,58])
    modelfit(X_train, y_train, X_test, y_test, document,  title= '2016 W/ previous vote', performCV=False, printFeatureImportance=True, cv_folds=5)
    ## Without Voted previous
    dat = dat.drop('voted_previous', axis = 1)
    X_train, X_test, y_train, y_test = prep_data(dat, cols = 2, cats=[52,53,54,55,56,57])
    modelfit(X_train, y_train, X_test, y_test, document,  title= '2016 W/out previous vote', performCV=False, printFeatureImportance=True, cv_folds=5)
    
    document.save('ERS_output.docx')
    print()
    print()
    print(f"Models completed and saved.")
    print()
    print()

main()
