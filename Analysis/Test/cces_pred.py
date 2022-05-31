import pandas as pd
import numpy as np
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.ensemble import VotingClassifier, AdaBoostClassifier, GradientBoostingClassifier, RandomForestClassifier
from sklearn.model_selection import GridSearchCV
from sklearn.decomposition import PCA
from sklearn.metrics import roc_curve, auc, confusion_matrix, accuracy_score, roc_auc_score
from sklearn.model_selection import train_test_split
from sklearn.model_selection import cross_val_score
import matplotlib.pylab as plt
#%matplotlib inline
from matplotlib.pylab import rcParams
rcParams['figure.figsize'] = 12, 4
from xgboost.sklearn import XGBClassifier
from docx import Document
from docx.text.paragraph import Paragraph
import os
from io import BytesIO
from tqdm import tqdm


def modelfit(alg, X_train, y_train, X_test, y_test, year, doc, performCV=True, printFeatureImportance=True, cv_folds=5):
    
    #alg.fit(X_train,y_train)

    y_preds = alg.predict(X_test)
    dtrain_predprob = alg.predict_proba(X_test)[:,1]
    
    #Print model report:
    doc.add_heading(f"\nModel Report {year}",0)
    doc.add_paragraph("Accuracy : %.4g" % accuracy_score(y_test, y_preds))
    doc.add_paragraph("AUC Score (Test): %f" % roc_auc_score(y_test, dtrain_predprob))

    #Perform cross-validation:
    if performCV:
        cv_score = cross_val_score(alg, X_train, y_train, cv=cv_folds, scoring='roc_auc')
        doc.add_paragraph("CV Score : Mean - %.7g | Std - %.7g | Min - %.7g | Max - %.7g" % (np.mean(cv_score),np.std(cv_score),np.min(cv_score),np.max(cv_score)))   
    #Print Feature Importance:
    if printFeatureImportance:
        feat_plot = BytesIO()
        feat_imp = pd.Series(alg.feat_importances_ure, X_train.columns).sort_values(ascending=False)
        feats = feat_imp.to_dict()
        for key in feats:
            doc.add_paragraph(f"{key}: {feats[key]}")
        feat_imp.plot(kind='bar', title='Feature Importances')
        plt.ylabel(f'Feature Importance Score for {year}')
        plt.savefig(feat_plot)

        doc.add_picture(feat_plot)
        feat_plot.close()

def data_prep(data, year):
    df = data[data.year == year]
    df = df.drop('year', axis = 1)
    if year <= 2012:
        df = df.drop('vv_party_prm', axis = 1)
        if year == 2006:
            df = df.drop('religion', axis=1)
    df = df.dropna()

    return df

def tune_model(X_train, y_train):
    param_test1 = {
 'max_depth':range(5,10,1),
 'min_child_weight':range(1,6,1)
    }
    gsearch1 = GridSearchCV(estimator = XGBClassifier( learning_rate =0.1, n_estimators=140, max_depth=5,
    min_child_weight=1, gamma=0, subsample=0.8, colsample_bytree=0.8,
    objective= 'binary:logistic', nthread = 6, scale_pos_weight=1, seed=27), 
    param_grid = param_test1, scoring='roc_auc',n_jobs=6, cv=5)
    gsearch1.fit(X_train,y_train)
    depth = gsearch1.best_params_['max_depth']
    child_weight = gsearch1.best_params_['min_child_weight']

    param_test2 = {
    'gamma':[i/10.0 for i in range(0,5)]
    }
    gsearch2 = GridSearchCV(estimator = XGBClassifier( learning_rate =0.1, n_estimators=140, max_depth=depth,
    min_child_weight=child_weight, gamma=0, subsample=0.8, colsample_bytree=0.8,
    objective= 'binary:logistic', nthread = 6, scale_pos_weight=1, seed=27), 
    param_grid = param_test2, scoring='roc_auc',n_jobs=6, cv=5)
    gsearch2.fit(X_train,y_train)
    gamm = gsearch2.best_params_['gamma']


    param_test3 = {
 'subsample':[i/10.0 for i in range(6,10)],
 'colsample_bytree':[i/10.0 for i in range(6,10)]
    }

    gsearch3 = GridSearchCV(estimator = XGBClassifier( learning_rate =0.1, n_estimators=140, max_depth=depth,
    min_child_weight=child_weight, gamma=gamm, subsample=0.8, colsample_bytree=0.8,
    objective= 'binary:logistic', nthread = 6, scale_pos_weight=1, seed=27), 
    param_grid = param_test3, scoring='roc_auc',n_jobs=6, cv=5)
    gsearch3.fit(X_train,y_train)
    sub_sample = gsearch3.best_params_['subsample']
    colsample = gsearch3.best_params_['colsample_bytree']

    param_test4 = {
 'reg_alpha':[0, 0.001, 0.005, 0.01, 0.05,1e-5, 1e-2, 0.1, 1, 100]
    }
    gsearch4 = GridSearchCV(estimator = XGBClassifier( learning_rate =0.1, n_estimators=140, max_depth=depth,
    min_child_weight=child_weight, gamma=gamm, subsample=sub_sample, colsample_bytree=colsample,
    objective= 'binary:logistic', nthread = 6, scale_pos_weight=1, seed=27), 
    param_grid = param_test4, scoring='roc_auc',n_jobs=6, cv=5)
    gsearch4.fit(X_train,y_train)  
    alpha = gsearch4.best_params_['reg_alpha']

    return depth, child_weight,gamm, sub_sample, colsample, alpha



def get_results(X_train, X_test, y_train, y_test, Year, doc, depth, child_weight, gamm, sub_sample, colsample, alpha, add_title = "", CV = False):
       
    xgb = XGBClassifier( learning_rate =0.01, n_estimators= 2000, max_depth=depth,
    min_child_weight=child_weight, gamma=gamm, subsample=sub_sample, colsample_bytree=colsample,
 objective= 'binary:logistic', nthread = 6, scale_pos_weight=1, seed=27, reg_alpha = alpha)
    
    xgb.fit(X_train,y_train)
    
    year  = str(Year) + add_title
    modelfit(xgb, X_train, y_train, X_test, y_test, year, doc, performCV =False)
    print()
    print()
    print(f"Results printed for {year}")
    print()
    print()

def do(data, year, doc):
    df = data_prep(data, year)
    y = df['voted']
    X = df.iloc[:, 2:]
    X_train, X_test, y_train, y_test = train_test_split(X,y, test_size= .25, random_state= 123)
    depth, child_weight, gamm, sub_sample, colsample, alpha = tune_model(X_train, y_train)
    get_results(X_train, X_test, y_train, y_test, year, doc, depth, child_weight,gamm, sub_sample, colsample, alpha,  CV = False)
    
    df = df.drop('citizen', axis = 1)
    y = df['voted']
    X = df.iloc[:, 2:]
    X_train, X_test, y_train, y_test = train_test_split(X,y, test_size= .25, random_state= 123)
    get_results(X_train, X_test, y_train, y_test, year, doc, depth, child_weight, gamm, sub_sample, colsample, alpha, add_title = " without citizenship", CV = False)
    
    doc.save('output.docx')


def main():
    document = Document()
    os.chdir(r'C:\Users\seanm\OneDrive\Documents\C - Research\NonVoters\Voter_Predictions\Paper Drafts')
    data = pd.read_csv(r'C:\Users\seanm\OneDrive\Documents\C - Research\NonVoters\Voter_Predictions\Data and Preperation\CCES data\clean_cces_data.csv')
    
    for year in range(2006,2019,2):
        do(data, year, document)

    ## run with party
    Year = "2012-2018"
    df = data.dropna()
    df = df.drop('citizen', axis = 1)
    y = df['voted']
    X = df.iloc[:, 2:]
    X_train, X_test, y_train, y_test = train_test_split(X,y, test_size= .25, random_state= 123)
    depth, child_weight, gamm, sub_sample, colsample, alpha = tune_model(X_train, y_train)
    get_results(X_train, X_test, y_train, y_test, Year, document, depth, child_weight, gamm, sub_sample, colsample, alpha,add_title = "'s elections with registered party", CV = False)
    document.save('output.docx')

    ## run without party
    df = df.drop('vv_party_prm', axis = 1)
    y = df['voted']
    X = df.iloc[:, 2:]
    X_train, X_test, y_train, y_test = train_test_split(X,y, test_size= .25, random_state= 123)
    depth, child_weight, gamm, sub_sample, colsample, alpha = tune_model(X_train, y_train)
    get_results(X_train, X_test, y_train, y_test, Year, document, depth, child_weight,gamm, sub_sample, colsample, alpha, add_title = "'s elections without registered party", CV = False)
    
    document.save('output.docx')
    print()
    print()
    print(f"Models completed and saved.")

main()